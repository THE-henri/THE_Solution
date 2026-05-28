"""
QY method text templates and Word document generator.

Each entry in METHOD_TEXTS is a plain-text method description suitable for
copying into a manuscript.  generate_case_method_doc() writes it to a
formatted .docx file using python-docx.
"""

from __future__ import annotations

from pathlib import Path

# ── Method text templates ─────────────────────────────────────────────────────
# Placeholders use [BRACKETS] for values the user fills in.

METHOD_TEXTS: dict[str, str] = {}

METHOD_TEXTS["A_only"] = """\
Quantum Yield Determination — A_only Case

Photochemical quantum yields Φ_AB for the A→B photoisomerisation were
determined from time-resolved UV/Vis absorbance data recorded at [MONITORING
WAVELENGTH(S)] nm.  Solutions of [COMPOUND] in [SOLVENT] at [TEMPERATURE] °C
were irradiated at [IRRADIATION WAVELENGTH] nm using a [LIGHT SOURCE]
(photon flux N = [VALUE] mol s⁻¹, determined by [ACTINOMETRY / POWER METER]).

Kinetic Model
The photochemical conversion of A to B was modelled assuming that only species
A absorbs at the irradiation wavelength (ε_B_irr ≈ 0) and that no thermal or
photochemical back-reaction occurs (k_th = 0, Φ_BA = 0).  Under these
conditions the time evolution of the concentration [A](t) is governed by:

    d[A]/dt = −(N/V) · l · F(t) · ε_A · Φ_AB · [A]

where N (mol s⁻¹) is the incident photon flux, V the irradiated volume, l the
optical path length (1 cm), ε_A the molar absorption coefficient of A at the
irradiation wavelength, and F(t) the photokinetic factor:

    F(t) = (1 − 10^(−A_irr(t))) / A_irr(t)

This factor corrects for the fraction of photons actually absorbed by the
sample at each time point, accounting for the decrease in absorbance as A is
consumed.  At low absorbance F → ln(10) ≈ 2.303.

Parameter Estimation
Quantum yields were extracted by numerical integration of the ODE followed by
nonlinear least-squares minimisation (Levenberg–Marquardt algorithm, lmfit
library) of the sum of squared residuals between the simulated and measured
absorbance traces at all monitoring wavelengths.  The initial concentration
[A]₀ was derived from the absorbance at the onset of irradiation and the
molar absorption coefficient ε_A([MONITORING WAVELENGTH]).

Uncertainty
The total uncertainty σ_total combines the fit standard error (σ_fit, from the
covariance matrix of the least-squares fit) and the photon flux uncertainty
(σ_I₀, from a perturbation analysis at N ± σ_N) in quadrature:

    σ_total = √(σ_fit² + σ_I₀²)

Results are reported as Φ_AB ± σ_total.
"""

METHOD_TEXTS["A_only_linearized"] = """\
Quantum Yield Determination — Exact Linearized Method (A_only Case)

Photochemical quantum yields Φ_AB were determined using the exact analytical
solution of the A_only photokinetic ODE.  Solutions of [COMPOUND] in [SOLVENT]
at [TEMPERATURE] °C were irradiated at [IRRADIATION WAVELENGTH] nm; kinetics
were monitored at [MONITORING WAVELENGTH] nm.

Analytical Solution
For the A_only case (ε_B = 0, Φ_BA = 0, k_th = 0), the variables in the ODE
separate exactly.  Integrating from t = 0 to t yields:

    log₁₀(10^{A_irr(t)} − 1) − log₁₀(10^{A_irr(0)} − 1)
        = −(N · ε_A_irr · l / V) · Φ_AB · t

where A_irr(t) = (ε_A_irr / ε_A_mon) · A_mon(t) is the absorbance at the
irradiation wavelength reconstructed from the monitored absorbance A_mon(t)
using the ratio of extinction coefficients.  This relationship is exact at all
absorbance levels (no dilute-solution approximation).

Linearization
Defining y(t) = log₁₀(10^{A_irr(t)} − 1), the equation becomes linear in t:

    y(t) = y(0) − slope · t
    slope = N · ε_A_irr · l · Φ_AB / V

A plot of y(t) versus t yields a straight line whose slope gives Φ_AB directly
without numerical integration.  The quantum yield was obtained by linear
regression over the irradiation window:

    Φ_AB = slope · V / (N · ε_A_irr · l)

Photon flux N was determined by [ACTINOMETRY / POWER METER].  Extinction
coefficients ε_A were determined from [SOURCE].
"""

METHOD_TEXTS["A_only_thermal"] = """\
Quantum Yield Determination — A_only_thermal Case

Photochemical quantum yields Φ_AB for the A→B photoisomerisation of the
thermally reversible (T-type) photoswitch [COMPOUND] were determined from
time-resolved UV/Vis absorbance data recorded at [MONITORING WAVELENGTH(S)] nm.
Solutions in [SOLVENT] at [TEMPERATURE] °C were irradiated at [IRRADIATION
WAVELENGTH] nm (photon flux N = [VALUE] mol s⁻¹).

Kinetic Model
Species A is the only light-absorbing isomer at the irradiation wavelength
(ε_B_irr ≈ 0).  Species B undergoes first-order thermal relaxation back to A
with rate constant k_th = [VALUE] s⁻¹ (half-life t½ = [VALUE] s at
[TEMPERATURE] °C, determined from independent thermal relaxation experiments).
The net rate equation is:

    d[A]/dt = −(N/V) · l · F(t) · ε_A · Φ_AB · [A] + k_th · [B]
    d[B]/dt = +(N/V) · l · F(t) · ε_A · Φ_AB · [A] − k_th · [B]

where F(t) = (1 − 10^(−ε_A · [A] · l)) / (ε_A · [A] · l) is the photokinetic
factor.  [A] + [B] = [A]₀ is conserved.

Parameter Estimation
The ODE was integrated numerically (scipy.integrate.odeint) and Φ_AB was
extracted by Levenberg–Marquardt least-squares minimisation.  The thermal rate
constant k_th was held fixed at its independently measured value; only Φ_AB was
optimised.

Uncertainty
σ_total = √(σ_fit² + σ_I₀²), where σ_fit is the fit covariance standard error
and σ_I₀ arises from perturbation of N by ±σ_N.
"""

METHOD_TEXTS["AB_both"] = """\
Quantum Yield Determination — AB_both Case

Forward and reverse quantum yields (Φ_AB, Φ_BA) for the photoisomerisation
A ⇌ B of [COMPOUND] were determined simultaneously from time-resolved UV/Vis
absorbance data.  Solutions in [SOLVENT] at [TEMPERATURE] °C were irradiated
at [IRRADIATION WAVELENGTH] nm (N = [VALUE] mol s⁻¹).

Kinetic Model
Both isomers A and B absorb at the irradiation wavelength.  The coupled ODEs are:

    d[A]/dt = (N/V) · l · F(t) · (ε_B · Φ_BA · [B] − ε_A · Φ_AB · [A])
    d[B]/dt = −d[A]/dt

where the photokinetic factor F(t) = (1 − 10^(−A_tot(t))) / A_tot(t) accounts
for the total absorbance A_tot = (ε_A · [A] + ε_B · [B]) · l of the mixture.

Parameter Estimation
Both Φ_AB and Φ_BA were fitted simultaneously by minimising the sum of squared
residuals between simulated and measured absorbance traces at all monitoring
wavelengths (Levenberg–Marquardt, lmfit).

Uncertainty
σ_total(Φ_AB) = √(σ_fit² + σ_I₀²); analogously for Φ_BA.
"""

METHOD_TEXTS["AB_thermal"] = """\
Quantum Yield Determination — AB_thermal Case

Forward and reverse quantum yields (Φ_AB, Φ_BA) and the thermal back-reaction
were jointly modelled for [COMPOUND] in [SOLVENT] at [TEMPERATURE] °C.
Irradiation was at [IRRADIATION WAVELENGTH] nm (N = [VALUE] mol s⁻¹).

Kinetic Model
Both isomers absorb at the irradiation wavelength, and species B undergoes
first-order thermal relaxation (k_th = [VALUE] s⁻¹):

    d[A]/dt = (N/V) · l · F(t) · (ε_B · Φ_BA · [B] − ε_A · Φ_AB · [A])
              + k_th · [B]
    d[B]/dt = −d[A]/dt

The photokinetic factor F(t) = (1 − 10^(−A_tot(t))) / A_tot(t) with
A_tot = (ε_A · [A] + ε_B · [B]) · l.  k_th was held fixed at its measured
value; Φ_AB and Φ_BA were the free parameters.

Parameter Estimation
Levenberg–Marquardt least-squares minimisation (lmfit) over all monitoring
wavelengths simultaneously.
"""

METHOD_TEXTS["A_thermal_PSS"] = """\
Quantum Yield Determination — PSS Algebraic Method (A_thermal_PSS)

The quantum yield Φ_AB for the photoisomerisation A→B of the T-type photoswitch
[COMPOUND] was determined from the photostationary state (PSS) reached under
continuous irradiation at [IRRADIATION WAVELENGTH] nm.

Principle
At PSS the net rate of formation of B equals zero.  For a system where only A
absorbs at the irradiation wavelength this gives the algebraic balance:

    Φ_AB · N · (1 − 10^(−A_PSS)) / V = k_th · [B]_PSS

Solving for Φ_AB:

    Φ_AB = k_th · [B]_PSS · V / (N · (1 − 10^(−A_PSS)))

where:
  k_th  = [VALUE] s⁻¹ — first-order thermal B→A rate constant at [TEMPERATURE] °C
           (determined from independent half-life measurements)
  [B]_PSS = concentration of B at PSS, derived from [PSS FRACTION / ABSORBANCE]
  A_PSS   = total absorbance at the irradiation wavelength at PSS
  N       = incident photon flux ([VALUE] mol s⁻¹, from [SOURCE])
  V       = sample volume ([VALUE] mL)

This method does not require ODE fitting; it is analytically exact under the
assumption that a true photostationary state has been reached.
"""


# ── Word document generator ───────────────────────────────────────────────────

def generate_case_method_doc(case: str, output_path: Path) -> Path:
    """
    Write a formatted Word document (.docx) containing the method text for
    the given case to *output_path*.  Returns the path to the saved file.

    Parameters
    ----------
    case        : one of the METHOD_TEXTS keys
    output_path : full path including filename (.docx)
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    text = METHOD_TEXTS.get(case)
    if text is None:
        raise ValueError(
            f"No method text for case '{case}'. "
            f"Available: {list(METHOD_TEXTS)}")

    doc = Document()

    # ── Styles ────────────────────────────────────────────────────────────────
    # Title style
    title_para = doc.add_heading(level=1)
    title_para.clear()
    run = title_para.add_run("Quantum Yield Analysis — Method Description")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()  # spacer

    # ── Parse and write the method text ───────────────────────────────────────
    lines = text.strip().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Blank line → paragraph break
        if not line.strip():
            i += 1
            continue

        # First line of each section is the section heading
        if (i == 0 or (i > 0 and not lines[i - 1].strip()
                       and not line.startswith(" ")
                       and not line.startswith("   ")
                       and not line.startswith("    "))):
            # Check if next non-empty line is indented (body text follows)
            is_heading = (i == 0 or (
                i + 1 < len(lines) and lines[i + 1].strip() and
                not lines[i + 1].startswith(" ")))
            # Simpler heuristic: a line that is ALL CAPS or ends with "Case"
            # and is not indented is a heading
            stripped = line.strip()
            if (stripped and not stripped.startswith("d[") and
                    not stripped.startswith("F(t)") and
                    not stripped.startswith("y(t)") and
                    not stripped.startswith("log") and
                    not stripped.startswith("slope") and
                    not stripped.startswith("σ") and
                    not stripped.startswith("Φ") and
                    (stripped.endswith("Case") or
                     stripped.endswith("Method") or
                     stripped.endswith("Estimation") or
                     stripped.endswith("Uncertainty") or
                     stripped.endswith("Linearization") or
                     stripped.endswith("Principle") or
                     stripped.endswith("Model"))):
                h = doc.add_heading(level=2)
                h.clear()
                r2 = h.add_run(stripped)
                r2.font.size = Pt(11)
                r2.font.bold = True
                r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
                i += 1
                continue

        # Equation lines (indented with 4+ spaces)
        if line.startswith("    "):
            eq_para = doc.add_paragraph()
            eq_para.paragraph_format.left_indent = Inches(0.5)
            r3 = eq_para.add_run(line.strip())
            r3.font.name = "Courier New"
            r3.font.size = Pt(10)
            i += 1
            continue

        # Normal body paragraph
        para = doc.add_paragraph(line.strip())
        para.paragraph_format.space_after = Pt(4)
        i += 1

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph(
        "Generated by THE Solution — QY Analysis Tool. "
        "Replace all [PLACEHOLDER] values with your experimental parameters.")
    note.runs[0].font.size = Pt(8)
    note.runs[0].font.italic = True
    note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
